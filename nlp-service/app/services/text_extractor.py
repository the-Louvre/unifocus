"""
文本提取服务
支持从HTML、PDF、Word文档中提取纯文本
"""
import re
from typing import Optional
from io import BytesIO

from bs4 import BeautifulSoup
from pdfminer.high_level import extract_text as extract_pdf_text
from pdfminer.layout import LAParams
from docx import Document
from loguru import logger


class TextExtractor:
    """文本提取器"""
    
    def __init__(self):
        """初始化文本提取器"""
        pass
    
    def extract_from_html(self, html_content: str) -> str:
        """
        从HTML中提取纯文本
        
        Args:
            html_content: HTML内容字符串
            
        Returns:
            提取的纯文本
        """
        try:
            # 使用BeautifulSoup解析HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除script和style标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 获取文本内容
            text = soup.get_text()
            
            # 清理文本：去除多余空白
            text = self._clean_text(text)
            
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from HTML: {e}")
            raise ValueError(f"HTML extraction failed: {str(e)}")
    
    def extract_from_pdf(self, pdf_bytes: bytes) -> str:
        """
        从PDF中提取文本
        
        Args:
            pdf_bytes: PDF文件的字节流
            
        Returns:
            提取的纯文本
        """
        try:
            # 使用pdfminer提取文本
            pdf_file = BytesIO(pdf_bytes)
            
            # 配置布局参数以优化中文提取
            laparams = LAParams(
                line_margin=0.5,
                word_margin=0.1,
                char_margin=2.0,
                boxes_flow=0.5,
            )
            
            text = extract_pdf_text(
                pdf_file,
                laparams=laparams,
                codec='utf-8'
            )
            
            # 清理文本
            text = self._clean_text(text)
            
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {e}")
            raise ValueError(f"PDF extraction failed: {str(e)}")
    
    def extract_from_docx(self, docx_bytes: bytes) -> str:
        """
        从Word文档中提取文本
        
        Args:
            docx_bytes: Word文档的字节流
            
        Returns:
            提取的纯文本
        """
        try:
            docx_file = BytesIO(docx_bytes)
            doc = Document(docx_file)
            
            # 提取所有段落的文本
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            
            text = '\n'.join(paragraphs)
            
            # 清理文本
            text = self._clean_text(text)
            
            return text
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError(f"DOCX extraction failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        清理文本：去除特殊字符和多余空白
        
        Args:
            text: 原始文本
            
        Returns:
            清理后的文本
        """
        # 去除控制字符（保留换行符和制表符）
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        
        # 将多个连续空白字符替换为单个空格
        text = re.sub(r'[ \t]+', ' ', text)
        
        # 将多个连续换行符替换为单个换行符
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        
        # 去除行首行尾空白
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # 去除文本首尾空白
        text = text.strip()
        
        return text

